from datetime import datetime
from pathlib import Path

from docx import Document


def generate_report(tender_json: dict, ranking: list[dict], reports_dir: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    report_path = reports_dir / f"{timestamp}-shortlist.docx"
    document = Document()
    document.add_heading("SmartTender AI - Shortlist Report", level=0)
    document.add_paragraph(f"Tender: {tender_json['title']}")
    document.add_paragraph(
        f"Required skills: {', '.join(tender_json.get('required_skills', [])) or 'Not detected'}"
    )
    document.add_paragraph(
        f"Required experience: {tender_json.get('experience_level', 0)} years"
    )

    for index, item in enumerate(ranking, start=1):
        document.add_heading(f"{index}. {item['candidate_name']} - {item['score']}/100", level=1)
        document.add_paragraph(f"Source file: {item['source_file']}")
        document.add_paragraph(f"Semantic score: {item['semantic_score']}/100")
        document.add_paragraph(f"Lexical score: {item['lexical_score']}/100")
        document.add_paragraph(
            f"Matched skills: {', '.join(item['matched_skills']) or 'No direct match detected'}"
        )
        document.add_paragraph(
            f"Missing skills: {', '.join(item['missing_skills']) or 'None'}"
        )
        for note in item["justification"]:
            document.add_paragraph(note, style="List Bullet")

    document.save(report_path)
    return report_path
